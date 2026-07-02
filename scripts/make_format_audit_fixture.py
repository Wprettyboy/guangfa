from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Mm, Pt


OUT = Path(r"C:\Users\23811\Desktop\广发new\格式审核测试样例-故意不规范.docx")


def set_run(run, font="宋体", size=12, bold=False):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold


def qn(tag):
    from docx.oxml.ns import qn as _qn

    return _qn(tag)


def add_body(doc, text, font="宋体", size=12, first_line=False, line_spacing=1.0, before=6, after=12):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(0 if not first_line else 32)
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    run = paragraph.add_run(text)
    set_run(run, font=font, size=size)
    return paragraph


def add_heading(doc, text, style, font="宋体", size=12, bold=False):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_run(run, font=font, size=size, bold=bold)
    return paragraph


def main():
    doc = Document()

    # Wrong margins: the app default expects 37/26/35/28 mm.
    section = doc.sections[0]
    section.top_margin = Mm(20)
    section.right_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(20)

    # Make Normal intentionally different from the configured standard.
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    # Static-looking TOC lines so the TOC rule has something to detect.
    for name in ["TOC 1", "TOC 2"]:
        if name not in [style.name for style in doc.styles]:
            doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph("目录", style="Title")
    doc.add_paragraph("第一章 项目概况........................................1", style="TOC 1")
    doc.add_paragraph("一、采购范围..........................................2", style="TOC 2")
    doc.add_page_break()

    # Missing heading style: looks like a chapter title, but uses Normal.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("第一章 项目概况")
    set_run(run, font="黑体", size=16, bold=True)

    # Body font/size/indent/line/spacing all intentionally wrong.
    add_body(
        doc,
        "本段正文故意使用宋体十二号、无首行缩进、单倍行距，并设置了异常段前段后，用于测试正文字体、字号、首行缩进、行距和段前段后的修复效果。",
        font="宋体",
        size=12,
        first_line=False,
        line_spacing=1.0,
        before=9,
        after=15,
    )
    add_body(
        doc,
        "本段正文故意使用微软雅黑十四号，和上一段不一致，用于测试批量统一正文字体字号时是否只修改被勾选的问题项。",
        font="微软雅黑",
        size=14,
        first_line=False,
        line_spacing=2.0,
        before=12,
        after=0,
    )

    # Consecutive blank lines.
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")

    # Body text mistakenly in Word outline.
    add_heading(
        doc,
        "这是一段很长的正文内容，故意套用了标题一样式，所以会误入 Word 导航窗格，修复后应该退出大纲。",
        "Heading 1",
        font="宋体",
        size=12,
        bold=False,
    )

    # Heading level mismatch: looks like level 2 but uses Heading 3.
    add_heading(doc, "一、采购范围", "Heading 3", font="宋体", size=12, bold=False)
    add_body(doc, "采购范围正文内容用于观察标题层级修复后，标题是否进入正确的大纲级别。", font="宋体", size=12)

    # Heading visual style issue: real headings with wrong fonts/sizes.
    add_heading(doc, "第二章 技术要求", "Heading 1", font="宋体", size=14, bold=False)
    add_heading(doc, "一、总体要求", "Heading 2", font="宋体", size=12, bold=False)
    add_heading(doc, "1. 详细要求", "Heading 3", font="宋体", size=11, bold=False)
    add_body(doc, "这些标题故意使用错误字体字号，用于测试标题字体字号配置项。", font="宋体", size=12)

    # Split heading: adjacent heading paragraphs that should be one title.
    add_heading(doc, "二、技术要求的", "Heading 2", font="宋体", size=12, bold=False)
    add_heading(doc, "总体说明", "Heading 2", font="宋体", size=12, bold=False)
    add_body(doc, "上面两个标题段落用于测试“标题拆分（合并被断开的标题段落）”。", font="宋体", size=12)

    # Add a second section with a different margin to trigger section normalization.
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.top_margin = Mm(15)
    new_section.right_margin = Mm(30)
    new_section.bottom_margin = Mm(25)
    new_section.left_margin = Mm(35)
    add_heading(doc, "第三章 其他说明", "Heading 1", font="宋体", size=13, bold=False)
    add_body(doc, "本页使用另一套页边距，用于测试页面版式中的页边距统一。", font="宋体", size=12)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
